"""
This code sample shows Prebuilt Read operations with the Azure Form Recognizer client library. 
The async versions of the samples require Python 3.6 or later.

To learn more, please visit the documentation - Quickstart: Document Intelligence (formerly Form Recognizer) SDKs
https://learn.microsoft.com/azure/ai-services/document-intelligence/quickstarts/get-started-sdks-rest-api?pivots=programming-language-python
"""

from azure.core.credentials import AzureKeyCredential
from azure.ai.formrecognizer import DocumentAnalysisClient
import docx
from fpdf import FPDF 
import comtypes.client 
import os

"""
Remember to remove the key from your code when you're done, and never post it publicly. For production, use
secure methods to store and access your credentials. For more information, see 
https://docs.microsoft.com/en-us/azure/cognitive-services/cognitive-services-security?tabs=command-line%2Ccsharp#environment-variables-and-application-configuration
"""
endpoint = "https://eastus.api.cognitive.microsoft.com/"
key = "fa301bd977644248bb7e0fd2046ad204"

def format_bounding_box(bounding_box):
    if not bounding_box:
        return "N/A"
    return ", ".join(["[{}, {}]".format(p.x, p.y) for p in bounding_box])

def convert_to_pdf_1(docx_filename, pdf_filename):
    # Simple conversion of text in DOCX to PDF using FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size = 12)
    doc = docx.Document(docx_filename)
    for para in doc.paragraphs:
        pdf.cell(200, 10, txt = para.text, ln = True)
    pdf.output(pdf_filename)

def convert_to_pdf(docx_filename, pdf_filename):
    # Use Microsoft Word to convert the Word document to a PDF
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(docx_filename)
    doc.SaveAs(pdf_filename, FileFormat=17)  # FileFormat=17 for PDF
    doc.Close()
    word.Quit()

def analyze_read():
    # sample document
    formUrl = "https://raw.githubusercontent.com/MalithSrineth/Document-Recognition/main/Assets/image02.jpg"

    document_analysis_client = DocumentAnalysisClient(
        endpoint=endpoint, credential=AzureKeyCredential(key)
    )
    
    poller = document_analysis_client.begin_analyze_document_from_url(
            "prebuilt-read", formUrl)
    result = poller.result()

    print ("Document contains content: ", result.content)

    base_dir = os.path.expanduser("~\\Documents\DXDY\Document Recognition")  # Change this to a suitable folder where you have write permissions
    if not os.path.exists(base_dir):
        os.makedirs(base_dir)

    if result.content:
        print("Document contains content")
        # Create a new Word document
        doc = docx.Document()
        doc.add_paragraph(result.content)
        # Save the Word document
        word_filename = os.path.join(base_dir, "output_2.docx")
        doc.save(word_filename)
        
        # Convert Word to PDF
        pdf_filename = os.path.join(base_dir, "output_2.pdf")
        convert_to_pdf(word_filename, pdf_filename)
    else:
        print("No content found in the document.")
    
    for idx, style in enumerate(result.styles):
        print(
            "Document contains {} content".format(
                "handwritten" if style.is_handwritten else "no handwritten"
            )
        )

    for page in result.pages:
        print("----Analyzing Read from page #{}----".format(page.page_number))
        print(
            "Page has width: {} and height: {}, measured with unit: {}".format(
                page.width, page.height, page.unit
            )
        )

        for line_idx, line in enumerate(page.lines):
            print(
                "...Line # {} has text content '{}' within bounding box '{}'".format(
                    line_idx,
                    line.content,
                    format_bounding_box(line.polygon),
                )
            )

        for word in page.words:
            print(
                "...Word '{}' has a confidence of {}".format(
                    word.content, word.confidence
                )
            )

    print("----------------------------------------")


if __name__ == "__main__":
    analyze_read()
